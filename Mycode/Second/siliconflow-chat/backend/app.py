import os
import re
import uuid
import pytesseract
import requests
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from PIL import Image
from dotenv import load_dotenv
import fitz  # PyMuPDF
import docx
import logging
import subprocess  # 用于调用外部工具
import tempfile  # 用于创建临时文件
import platform  # 用于检测操作系统

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# 加载环境变量
load_dotenv()

# 创建Flask应用
app = Flask(__name__)
CORS(app)

# 配置API密钥和模型映射
API_KEY = os.getenv("QWEN_API_KEY")
if not API_KEY:
    logger.error("QWEN_API_KEY environment variable is not set.")
    raise ValueError("QWEN_API_KEY environment variable is not set.")

BASE_URL = "https://api.siliconflow.cn/v1"
MODEL_MAP = {
    "qwen": "Qwen/Qwen3-32B",
    "deepseek": "deepseek-ai/DeepSeek-R1",
    "qvq": "Qwen/QVQ-72B-Preview"
}

# 文件存储配置
UPLOAD_FOLDER = os.path.join(os.getcwd(), 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# 存储文件元数据 {file_id: {filename, file_path}}
file_metadata = {}

# 允许的文件类型
ALLOWED_EXTENSIONS = {'jpg', 'jpeg', 'png', 'gif', 'txt', 'tex', 'doc', 'docx', 'pdf'}

# 设置文件大小限制 (10MB)
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024

def allowed_file(filename):
    """检查文件扩展名是否允许"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def ocr_image(file_path):
    """使用OCR识别图像中的文本"""
    try:
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image, lang='eng+chi_sim')
        return text.strip()
    except Exception as e:
        logger.error(f"OCR失败: {str(e)}")
        return f"图像OCR失败: {str(e)}"

def extract_text_from_pdf(file_path):
    """提取PDF文本内容"""
    try:
        text = ""
        with fitz.open(file_path) as doc:
            for page in doc:
                text += page.get_text()
        return text.strip()
    except Exception as e:
        logger.error(f"PDF处理失败: {str(e)}")
        return f"PDF处理失败: {str(e)}"

def extract_text_from_docx(file_path):
    """提取DOCX文本内容"""
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # 提取段落文本
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        full_text.append(para.text)
        
        return "\n".join(full_text).strip()
    except Exception as e:
        logger.error(f"DOCX处理失败: {str(e)}")
        return f"DOCX处理失败: {str(e)}"

def convert_doc_to_docx(doc_path):
    """将DOC文件转换为DOCX格式"""
    try:
        # 创建临时文件
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            docx_path = temp_file.name
        
        # 根据操作系统使用不同的转换工具
        if platform.system() == 'Windows':
            # Windows: 使用Word COM对象
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(doc_path)
            doc.SaveAs2(docx_path, FileFormat=16)  # 16 = wdFormatDocumentDefault (DOCX)
            doc.Close()
            word.Quit()
        else:
            # macOS/Linux: 使用LibreOffice
            cmd = ['libreoffice', '--headless', '--convert-to', 'docx', '--outdir', 
                   os.path.dirname(docx_path), doc_path]
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        
        return docx_path
    except Exception as e:
        logger.error(f"DOC转DOCX失败: {str(e)}")
        return None

def extract_text_from_doc(file_path):
    """提取DOC文本内容 - 通过转换为DOCX"""
    try:
        # 首先尝试转换为DOCX
        docx_path = convert_doc_to_docx(file_path)
        if not docx_path or not os.path.exists(docx_path):
            return "无法转换DOC文件为DOCX格式"
        
        # 然后提取转换后的DOCX内容
        text = extract_text_from_docx(docx_path)
        
        # 删除临时文件
        try:
            os.unlink(docx_path)
        except:
            pass
            
        return text
    except Exception as e:
        logger.error(f"DOC处理失败: {str(e)}")
        return f"DOC处理失败: {str(e)}"

def process_file(file_path, filename):
    """根据文件类型处理文件内容"""
    ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    
    if ext in ['jpg', 'jpeg', 'png', 'gif']:
        return ocr_image(file_path)
    elif ext == 'pdf':
        return extract_text_from_pdf(file_path)
    elif ext == 'docx':
        return extract_text_from_docx(file_path)
    elif ext == 'doc':
        return extract_text_from_doc(file_path)
    elif ext in ['txt', 'tex']:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read().strip()
        except Exception as e:
            logger.error(f"文本处理失败: {str(e)}")
            return f"文本处理失败: {str(e)}"
    else:
        return f"文件 {filename} 已上传，但当前仅支持图片/文档内容分析"

# @app.route("/chat", methods=["POST"])
# def chat():
#     """处理聊天请求 - 支持文件引用"""
#     data = request.get_json()
#     if not data:
#         return jsonify({"error": "请求数据为空"}), 400
        
#     model_alias = data.get("model", "deepseek").lower()
#     user_message = data.get("message", "")
#     history = data.get("history", [])

#     model = MODEL_MAP.get(model_alias, "deepseek-ai/DeepSeek-R1")

#     headers = {
#         "Authorization": f"Bearer {API_KEY}",
#         "Content-Type": "application/json"
#     }

#     # 1. 检查用户消息中是否包含文件名引用
#     file_contents = []
#     # 修复：匹配带空格的文件名和正确的扩展名
#     analysis_pattern = r'分析\s*([\w\s\-\.\u4e00-\u9fa5]+\.(?:jpg|jpeg|png|gif|txt|tex|docx?|pdf))'
#     analysis_match = re.search(analysis_pattern, user_message, re.IGNORECASE)
    
#     if analysis_match:
#         filename = analysis_match.group(1).strip()
#         logger.info(f"提取到文件名: {filename}")
        
#         # 2. 在已上传文件中查找（不区分大小写）
#         file_found = False
#         for file_id, meta in file_metadata.items():
#             # 修复：使用文件名部分匹配（忽略大小写）
#             if meta['filename'].lower() == filename.lower() or \
#                meta['filename'].lower().startswith(filename.split('.')[0].lower()):
#                 file_found = True
#                 try:
#                     # 处理文件内容
#                     content = process_file(meta['file_path'], meta['filename'])
#                     logger.info(f"文件内容 (前100字符): {content[:100]}...")  # 记录前100字符用于调试
#                     file_contents.append(f"文件 '{meta['filename']}' 的内容:\n{content}")
#                     break
#                 except Exception as e:
#                     file_contents.append(f"处理文件出错: {str(e)}")
#                     break
        
#         if not file_found:
#             available_files = [meta['filename'] for meta in file_metadata.values()]
#             logger.warning(f"未找到文件 '{filename}'，已上传文件: {available_files}")
#             file_contents.append(f"未找到文件 '{filename}'。已上传文件: {', '.join(available_files)}")
#     else:
#         logger.info("未检测到有效的文件分析指令")
#         file_contents.append("未检测到有效的文件分析指令。请使用格式: 分析文件名")

#     # 3. 构建系统消息（包含文件内容）
#     system_message = ""
#     if file_contents:
#         system_message = "\n\n".join(file_contents)
#         logger.info(f"系统消息内容: {system_message[:200]}...")  # 只打印前200字符
    
#     # 4. 构建消息历史
#     messages = []
#     if system_message:
#         # 添加文件内容作为系统消息
#         messages.append({"role": "system", "content": system_message})
    
#     # 添加历史消息
#     for msg in history:
#         messages.append({
#             "role": "user" if msg['from'] == 'user' else "assistant",
#             "content": msg['text']
#         })
    
#     # 添加当前用户消息
#     messages.append({"role": "user", "content": user_message})

#     # 5. 发送到API
#     payload = {
#         "model": model,
#         "messages": messages
#     }

#     try:
#         logger.info(f"发送请求到API: {model}")
#         res = requests.post(f"{BASE_URL}/chat/completions", headers=headers, json=payload)
#         res.raise_for_status()
#         resp_json = res.json()

#         reply = ""
#         if "choices" in resp_json and len(resp_json["choices"]) > 0:
#             reply = resp_json["choices"][0]["message"].get("content", "")
#         else:
#             reply = "未收到有效回复"

#         return jsonify({
#             "reply": reply,
#             "file_contents": file_contents  # 返回文件内容用于前端显示
#         })

#     except requests.exceptions.RequestException as e:
#         logger.error(f"请求API失败: {str(e)}")
#         return jsonify({"error": f"请求API失败: {str(e)}"}), 500
#     except Exception as e:
#         logger.error(f"服务器错误: {str(e)}")
#         return jsonify({"error": f"服务器错误: {str(e)}"}), 500
@app.route("/chat", methods=["POST"])
def chat():
    """处理聊天请求 - 支持文件引用"""
    data = request.get_json()
    if not data:
        return jsonify({"error": "请求数据为空"}), 400
        
    model_alias = data.get("model", "deepseek").lower()
    user_message = data.get("message", "")
    history = data.get("history", [])

    model = MODEL_MAP.get(model_alias, "deepseek-ai/DeepSeek-R1")

    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }

    # 1. 检查用户消息中是否包含文件名引用
    file_contents = []
    # 修复：匹配带空格的文件名和正确的扩展名
    analysis_pattern = r'分析\s*([\w\s\-\.\u4e00-\u9fa5]+\.(?:jpg|jpeg|png|gif|txt|tex|docx?|pdf))'
    analysis_match = re.search(analysis_pattern, user_message, re.IGNORECASE)
    
    if analysis_match:
        filename = analysis_match.group(1).strip()
        logger.info(f"提取到文件名: {filename}")
        
        # 2. 在已上传文件中查找（不区分大小写）
        file_found = False
        for file_id, meta in file_metadata.items():
            # 修复：使用文件名部分匹配（忽略大小写）
            if meta['filename'].lower() == filename.lower() or \
               meta['filename'].lower().startswith(filename.split('.')[0].lower()):
                file_found = True
                try:
                    # 处理文件内容
                    content = process_file(meta['file_path'], meta['filename'])
                    logger.info(f"文件内容 (前100字符): {content[:100]}...")  # 记录前100字符用于调试
                    file_contents.append(f"文件 '{meta['filename']}' 的内容:\n{content}")
                    break
                except Exception as e:
                    file_contents.append(f"处理文件出错: {str(e)}")
                    break
        
        if not file_found:
            available_files = [meta['filename'] for meta in file_metadata.values()]
            logger.warning(f"未找到文件 '{filename}'，已上传文件: {available_files}")
            file_contents.append(f"未找到文件 '{filename}'。已上传文件: {', '.join(available_files)}")
    # 移除未检测到文件分析指令的提示

    # 3. 构建系统消息（包含文件内容）
    system_message = ""
    if file_contents:
        system_message = "\n\n".join(file_contents)
        logger.info(f"系统消息内容: {system_message[:200]}...")  # 只打印前200字符
    
    # 4. 构建消息历史
    messages = []
    if system_message:
        # 添加文件内容作为系统消息
        messages.append({"role": "system", "content": system_message})
    
    # 添加历史消息
    for msg in history:
        messages.append({
            "role": "user" if msg['from'] == 'user' else "assistant",
            "content": msg['text']
        })
    
    # 添加当前用户消息
    messages.append({"role": "user", "content": user_message})

    # 5. 发送到API
    payload = {
        "model": model,
        "messages": messages
    }

    try:
        logger.info(f"发送请求到API: {model}")
        res = requests.post(f"{BASE_URL}/chat/completions", headers=headers, json=payload)
        res.raise_for_status()
        resp_json = res.json()

        reply = ""
        if "choices" in resp_json and len(resp_json["choices"]) > 0:
            reply = resp_json["choices"][0]["message"].get("content", "")
        else:
            reply = "未收到有效回复"

        return jsonify({
            "reply": reply,
            "file_contents": file_contents  # 返回文件内容用于前端显示
        })

    except requests.exceptions.RequestException as e:
        logger.error(f"请求API失败: {str(e)}")
        return jsonify({"error": f"请求API失败: {str(e)}"}), 500
    except Exception as e:
        logger.error(f"服务器错误: {str(e)}")
        return jsonify({"error": f"服务器错误: {str(e)}"}), 500


@app.route('/upload', methods=['POST'])
def upload_file():
    """处理文件上传"""
    if 'file' not in request.files:
        return jsonify({'error': '没有上传文件'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '未选择文件'}), 400

    if not allowed_file(file.filename):
        return jsonify({'error': '文件类型不允许'}), 400

    # 生成唯一文件名 - 保留原始文件名
    file_id = str(uuid.uuid4())
    original_name = file.filename  # 直接使用原始文件名
    
    # 使用UUID作为存储文件名，保留原始扩展名
    if '.' in original_name:
        ext = original_name.rsplit('.', 1)[1].lower()
        stored_name = f"{file_id}.{ext}"
    else:
        stored_name = file_id
    
    file_path = os.path.join(UPLOAD_FOLDER, stored_name)
    
    try:
        file.save(file_path)
    except Exception as e:
        logger.error(f"保存文件失败: {str(e)}")
        return jsonify({'error': f'保存文件失败: {str(e)}'}), 500

    # 存储文件元数据
    file_metadata[file_id] = {
        'filename': original_name,
        'file_path': file_path
    }

    logger.info(f"文件上传成功: {original_name} (ID: {file_id})")
    # 返回成功消息
    return jsonify({
        'file_id': file_id,
        'filename': original_name,
        'result': f"文件 {original_name} 上传成功"
    })

@app.route('/delete/<file_id>', methods=['DELETE'])
def delete_file(file_id):
    """删除上传的文件"""
    if file_id in file_metadata:
        try:
            file_path = file_metadata[file_id]['file_path']
            os.remove(file_path)
            del file_metadata[file_id]
            logger.info(f"文件删除成功: {file_id}")
            return jsonify({'success': True})
        except Exception as e:
            logger.error(f'删除文件失败: {str(e)}')
            return jsonify({'error': f'删除文件失败: {str(e)}'}), 500
    
    return jsonify({'error': '文件不存在'}), 404

@app.route('/files', methods=['GET'])
def list_files():
    """获取已上传文件列表"""
    files = []
    for file_id, meta in file_metadata.items():
        files.append({
            'file_id': file_id,
            'filename': meta['filename']
        })
    return jsonify(files)

if __name__ == "__main__":
    app.run(debug=True, port=5001)
